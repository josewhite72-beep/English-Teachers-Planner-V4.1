"""
Vercel Serverless Function - Export Planner to Word
Usando Flask (WSGI) - Más robusto que BaseHTTPRequestHandler
"""

from flask import Flask, request, Response
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import json

app = Flask(__name__)

@app.route('/api/export-docx', methods=['POST', 'OPTIONS'])
def export_docx():
    # Manejar preflight CORS
    if request.method == 'OPTIONS':
        response = Response()
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return response
    
    try:
        print("🚀 Función export-docx iniciada", flush=True)
        
        # Obtener datos del request
        planner_data = request.get_json()
        
        if not planner_data:
            print("❌ No se recibieron datos", flush=True)
            return Response(
                json.dumps({'error': 'No data provided'}), 
                status=400, 
                mimetype='application/json'
            )
        
        print(f"✅ Datos recibidos - Grade: {planner_data.get('grade')}", flush=True)
        
        # Generar documento
        doc = create_planner_document(planner_data)
        
        # Guardar en memoria
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        print(f"✅ Documento generado - Tamaño: {len(doc_io.getvalue())} bytes", flush=True)
        
        # Crear respuesta
        response = Response(
            response=doc_io.getvalue(),
            status=200,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response.headers['Content-Disposition'] = f'attachment; filename="planner_{planner_data.get("grade", "unknown")}.docx"'
        response.headers['Access-Control-Allow-Origin'] = '*'
        
        print("✅ Respuesta enviada exitosamente", flush=True)
        return response

    except Exception as e:
        print(f"❌ Error: {str(e)}", flush=True)
        import traceback
        traceback.print_exc()
        
        return Response(
            json.dumps({'error': str(e)}), 
            status=500, 
            mimetype='application/json'
        )


def create_planner_document(data):
    """Crear documento Word completo"""
    print("  → Creando documento...", flush=True)
    doc = Document()
    
    # Header MEDUCA
    add_meduca_header(doc)
    
    # Theme Planner
    add_theme_planner(doc, data)
    
    # Lesson Planners
    add_lesson_planners(doc, data)
    
    print("  ✓ Documento creado", flush=True)
    return doc


def add_meduca_header(doc):
    """Header oficial MEDUCA"""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = heading.add_run('MINISTERIO DE EDUCACIÓN\n')
    run1.bold = True
    run1.font.size = Pt(14)
    
    run2 = heading.add_run('Dirección Nacional de Currículo de Lengua Extranjera')
    run2.font.size = Pt(11)
    
    doc.add_paragraph()


def add_theme_planner(doc, data):
    """Theme Planner completo"""
    theme = data.get('theme_planner', {})
    
    # Título
    title = doc.add_heading('Theme Planner - Overview', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. General Information
    doc.add_heading('1. General Information:', level=2)
    add_general_info_table(doc, theme.get('general_information', {}))
    
    # 2. Standards
    doc.add_heading('2. Specific Standards and Learning Outcomes:', level=2)
    add_standards_table(doc, theme.get('standards_and_learning_outcomes', {}))
    
    # 3. Competences
    doc.add_heading('3. Communicative Competences', level=2)
    add_competences_table(doc, theme.get('communicative_competences', {}))
    
    # Project
    if 'project' in data:
        add_project_section(doc, data['project'])
    
    # 4. Objectives
    doc.add_heading('4. Specific Objectives', level=2)
    add_objectives_table(doc, theme.get('specific_objectives', {}))
    
    # 5. Materials
    doc.add_heading('5. Materials and Teaching Strategies', level=2)
    add_materials_table(doc, theme.get('materials_and_strategies', {}))
    
    # 6. Learning Sequence
    doc.add_heading('6. Learning Sequence', level=2)
    add_learning_sequence_table(doc, data.get('lesson_planners', []), theme.get('learning_sequence', {}))


def add_general_info_table(doc, general):
    """Tabla de información general"""
    grade_labels = {
        'pre_k': 'Pre-K', 'K': 'Kindergarten',
        '1': 'Grade 1', '2': 'Grade 2', '3': 'Grade 3', '4': 'Grade 4',
        '5': 'Grade 5', '6': 'Grade 6', '7': 'Grade 7', '8': 'Grade 8',
        '9': 'Grade 9', '10': 'Grade 10', '11': 'Grade 11', '12': 'Grade 12',
    }
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = '1. Teacher(s):'
    table.cell(0, 1).merge(table.cell(0, 3)).text = general.get('teachers', '_______________')
    
    table.cell(1, 0).text = '2. Grade:'
    table.cell(1, 1).text = grade_labels.get(general.get('grade', ''), general.get('grade', ''))
    table.cell(1, 2).text = '3. CEFR Level:'
    table.cell(1, 3).text = general.get('cefr_level', '______')
    
    table.cell(2, 0).text = '4. Trimester:'
    table.cell(2, 1).text = general.get('trimester', '______')
    table.cell(2, 2).text = '5. Weekly Hour(s):'
    table.cell(2, 3).text = general.get('weekly_hours', '______')
    
    table.cell(3, 0).text = '6. Week(s):'
    table.cell(3, 1).merge(table.cell(3, 3)).text = general.get('week_range', 'From week ____ to week ____')
    
    table.cell(4, 0).text = '7. Scenario __:'
    table.cell(4, 1).text = general.get('scenario', '')
    table.cell(4, 2).text = '8. Theme __:'
    table.cell(4, 3).text = general.get('theme', '')
    
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()


def add_standards_table(doc, standards):
    """Tabla de standards"""
    skills = ['listening', 'reading', 'speaking', 'writing', 'mediation']
    labels = ['1. Listening:', '2. Reading:', '3. Speaking:', '4. Writing:', '5. Mediation:']
    
    table = doc.add_table(rows=len(skills) + 1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Skills:'
    hdr_cells[1].text = 'Specific Standards:'
    hdr_cells[2].text = 'Learning Outcomes:'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    for idx, skill in enumerate(skills):
        row = table.rows[idx + 1]
        row.cells[0].text = labels[idx]
        row.cells[0].paragraphs[0].runs[0].bold = True
        
        skill_data = standards.get(skill, {})
        
        std_list = []
        if isinstance(skill_data.get('specific'), list):
            std_list = skill_data['specific']
        elif isinstance(skill_data.get('specific_standards'), list):
            std_list = skill_data['specific_standards']
        
        row.cells[1].text = '\n'.join([f"• {s}" for s in std_list]) if std_list else 'To be defined'
        
        outcomes = skill_data.get('learning_outcomes', [])
        row.cells[2].text = '\n'.join([f"• {o}" for o in outcomes]) if outcomes else 'To be defined'
    
    doc.add_paragraph()


def add_competences_table(doc, competences):
    """Tabla de competencias"""
    ling = competences.get('linguistic', {})
    prag = competences.get('pragmatic', '')
    socio = competences.get('sociolinguistic', '')
    
    grammar = ling.get('grammatical_features') or ling.get('grammar', [])
    vocab = ling.get('vocabulary', [])
    pronun = ling.get('phonemic_awareness') or ling.get('pronunciation', '')
    
    grammar_text = ', '.join(grammar) if isinstance(grammar, list) else str(grammar)
    
    # Manejar vocabulario como objeto o array
    if isinstance(vocab, list):
        vocab_text = ', '.join(vocab[:10])
    elif isinstance(vocab, dict):
        vocab_words = []
        for key in ['nouns', 'verbs', 'adjectives', 'prepositions']:
            if isinstance(vocab.get(key), list):
                vocab_words.extend(vocab[key])
        vocab_text = ', '.join(vocab_words[:10])
    else:
        vocab_text = str(vocab)
    
    prag_text = ', '.join(prag) if isinstance(prag, list) else str(prag)
    socio_text = ', '.join(socio) if isinstance(socio, list) else str(socio)
    
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Linguistic Competence\n(Learn to Know)'
    hdr[1].text = 'Pragmatic Competence\n(Learn to Do)'
    hdr[2].text = 'Sociolinguistic Competence\n(Learn to Be)'
    for cell in hdr:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
    
    cells = table.rows[1].cells
    ling_text = f"• Grammatical Features:\n{grammar_text or '______'}\n\n"
    ling_text += f"• Vocabulary:\n{vocab_text or '______'}\n\n"
    ling_text += f"• Pronunciation:\n{pronun or '______'}"
    cells[0].text = ling_text
    cells[1].text = prag_text or '______'
    cells[2].text = socio_text or '______'
    
    doc.add_paragraph()


def add_project_section(doc, project):
    """Sección de proyecto"""
    p = doc.add_paragraph()
    p.add_run('21st-Century Skills Project\n').bold = True
    p.add_run(f"Name: {project.get('name', '')}\n")
    p.add_run(f"Category: {project.get('category', '')}\n")
    p.add_run(f"Overview: {project.get('overview', '')}")
    doc.add_paragraph()


def add_objectives_table(doc, objectives):
    """Tabla de objetivos"""
    skills = ['listening', 'reading', 'speaking', 'writing', 'mediation']
    labels = ['Listening', 'Reading', 'Speaking', 'Writing', 'Mediation']
    
    table = doc.add_table(rows=len(skills), cols=2)
    table.style = 'Table Grid'
    
    for idx, skill in enumerate(skills):
        row = table.rows[idx]
        row.cells[0].text = labels[idx]
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = objectives.get(skill, 'To be completed')
    
    doc.add_paragraph()


def add_materials_table(doc, materials):
    """Tabla de materiales"""
    mat_list = materials.get('required_materials', [])
    diff = materials.get('differentiated_instruction', 'To be completed')
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = 'Materials'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = '\n'.join([f"• {m}" for m in mat_list]) if mat_list else 'To be completed'
    
    table.cell(1, 0).text = 'Differentiated Instruction and Accommodations for Students with Diverse Learning Needs (DLN)'
    table.cell(1, 0).paragraphs[0].runs[0].bold = True
    table.cell(1, 1).text = diff
    
    doc.add_paragraph()


def add_learning_sequence_table(doc, lessons, learning_seq):
    """Tabla de secuencia"""
    dates = learning_seq.get('lesson_dates', ['', '', '', '', ''])
    
    table = doc.add_table(rows=len(lessons) + 1, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Lesson'
    hdr[1].text = 'Learning Sequence'
    hdr[2].text = 'Date'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    
    for idx, lesson in enumerate(lessons):
        row = table.rows[idx + 1]
        row.cells[0].text = f"Lesson {lesson.get('lesson_number', idx+1)}"
        row.cells[0].paragraphs[0].runs[0].bold = True
        
        seq_text = f"{lesson.get('skill_focus', '')}\n{lesson.get('specific_objective', '')}"
        row.cells[1].text = seq_text
        row.cells[2].text = dates[idx] if idx < len(dates) else '______'
    
    doc.add_paragraph()


def add_lesson_planners(doc, data):
    """Agregar lesson planners"""
    lessons = data.get('lesson_planners', [])
    
    for idx, lesson in enumerate(lessons):
        if idx > 0:
            doc.add_page_break()
        
        title = doc.add_heading(f'Lesson Planner - Lesson #{lesson.get("lesson_number", idx+1)}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info table (simplificada para reducir complejidad)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = 'Lesson:'
        table.cell(0, 1).text = f"Lesson {lesson.get('lesson_number', idx+1)} - {lesson.get('skill_focus', '')}"
        
        table.cell(1, 0).text = 'Specific Objective:'
        table.cell(1, 1).text = lesson.get('specific_objective', '______')
        
        table.cell(2, 0).text = 'Learning Outcome:'
        table.cell(2, 1).text = lesson.get('learning_outcome', '______')
        
        doc.add_paragraph()
        
        # Stages
        doc.add_heading('The Six Lesson Stages', level=2)
        stages = lesson.get('lesson_stages', [])
        
        if stages:
            for stage in stages:
                activities = stage.get('activities', [])
                p = doc.add_paragraph()
                p.add_run(f"{stage.get('stage', 'Stage')} ({stage.get('estimated_time', '___')})\n").bold = True
                for act in activities:
                    p.add_run(f"• {act}\n")


# Vercel busca 'app' para WSGI
if __name__ == '__main__':
    app.run()
