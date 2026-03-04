from fastapi import FastAPI, APIRouter, HTTPException
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import json
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Dict, Any, Optional
import uuid
from datetime import datetime, timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import tempfile
from emergentintegrations.llm.chat import LlmChat, UserMessage

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Data paths
DATA_DIR = ROOT_DIR / 'data'
GRADES_DIR = DATA_DIR / 'grades'
PROJECTS_OFFICIAL_DIR = DATA_DIR / 'projects' / 'official'
PROJECTS_C21_DIR = DATA_DIR / 'projects' / 'c21'
INSTITUTIONAL_STANDARDS_DIR = DATA_DIR / 'institutional_standards'
SCOPE_SEQUENCE_FILE = DATA_DIR / 'scope_sequence.json'
GENERATED_PLANNERS_DIR = DATA_DIR / 'generated_planners'

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Models
class PlannerRequest(BaseModel):
    grade: str
    scenario: str
    theme: str
    plan_type: str  # basic, standard, enriched
    official_format: bool = False
    project_id: Optional[str] = None
    language: str = "es"
    # New MEDUCA official format fields
    teacher_name: Optional[str] = ""
    trimester: Optional[str] = ""
    weekly_hours: Optional[str] = ""
    week_from: Optional[str] = ""
    week_to: Optional[str] = ""
    
class InstitutionalData(BaseModel):
    school_name: str = ""
    teacher_name: str = ""
    trimester: str = ""
    week_range: str = ""

# Helper functions to load JSON data
def load_grade_data(grade: str) -> dict:
    """Load grade curriculum data"""
    file_path = GRADES_DIR / f"{grade}.json"
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Grade {grade} not found")
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_projects_official(grade: str) -> dict:
    """Load official projects for a grade"""
    file_path = PROJECTS_OFFICIAL_DIR / f"{grade}.json"
    if not file_path.exists():
        return {}
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_projects_c21(grade: str) -> dict:
    """Load C21 complementary projects for a grade"""
    file_path = PROJECTS_C21_DIR / f"{grade}.json"
    if not file_path.exists():
        return {}
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_institutional_standards(grade: str) -> dict:
    """Load institutional standards for a grade"""
    file_path = INSTITUTIONAL_STANDARDS_DIR / f"{grade}.json"
    if not file_path.exists():
        return {}
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_pregenerated_planner(grade: str, scenario: str, theme: str, plan_type: str, language: str) -> dict:
    """Load a pregenerated planner from JSON file if it exists"""
    # Extract scenario number from scenario string (e.g., "Scenario 1: All Week Long!" -> "1")
    import re
    scenario_match = re.search(r'Scenario (\d+)', scenario)
    scenario_num = scenario_match.group(1) if scenario_match else "1"
    
    # Determine theme number based on the theme list for that scenario
    # For now, we'll use a simple approach - if we can load the grade data, we can find the index
    try:
        grade_data = load_grade_data(grade)
        scenario_data = None
        if isinstance(grade_data, list):
            scenario_data = next((s for s in grade_data if s.get('scenario', s.get('title', '')) == scenario), None)
        elif isinstance(grade_data, dict) and 'scenarios' in grade_data:
            scenario_data = next((s for s in grade_data['scenarios'] if s.get('scenario', s.get('title', '')) == scenario), None)
        
        if scenario_data:
            themes = scenario_data.get('themes', [])
            theme_num = str(themes.index(theme) + 1) if theme in themes else "1"
        else:
            theme_num = "1"
    except:
        theme_num = "1"
    
    # Build filename: {grade}_{scenario_num}_{theme_num}_{plan_type}_{language}.json
    filename = f"{grade}_{scenario_num}_{theme_num}_{plan_type}_{language}.json"
    file_path = GENERATED_PLANNERS_DIR / filename
    
    logger.info(f"Looking for pregenerated planner: {filename}")
    
    if file_path.exists():
        logger.info(f"Found pregenerated planner: {filename}")
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    return None

# API Routes
@api_router.get("/")
async def root():
    return {"message": "English Teachers Planner API"}

@api_router.get("/grades")
async def get_grades():
    """Get list of available grades"""
    grades = ["pre_k", "K", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12"]
    return {"grades": grades}

@api_router.get("/grades/{grade}/scenarios")
async def get_scenarios(grade: str):
    """Get scenarios for a specific grade"""
    try:
        data = load_grade_data(grade)
        
        # Check if data is array or object with scenarios
        if isinstance(data, list):
            scenarios = [s.get('scenario', s.get('title', '')) for s in data]
        elif isinstance(data, dict) and 'scenarios' in data:
            scenarios = [s.get('scenario', s.get('title', '')) for s in data['scenarios']]
        else:
            scenarios = []
            
        return {"scenarios": scenarios}
    except Exception as e:
        logger.error(f"Error loading scenarios for grade {grade}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/grades/{grade}/scenarios/{scenario}/themes")
async def get_themes(grade: str, scenario: str):
    """Get themes for a specific scenario"""
    try:
        data = load_grade_data(grade)
        
        # Find the scenario
        scenario_data = None
        if isinstance(data, list):
            scenario_data = next((s for s in data if s.get('scenario', s.get('title', '')) == scenario), None)
        elif isinstance(data, dict) and 'scenarios' in data:
            scenario_data = next((s for s in data['scenarios'] if s.get('scenario', s.get('title', '')) == scenario), None)
            
        if not scenario_data:
            return {"themes": []}
            
        themes = scenario_data.get('themes', [])
        return {"themes": themes}
    except Exception as e:
        logger.error(f"Error loading themes: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/grades/{grade}/content")
async def get_grade_content(grade: str):
    """Get complete content for a grade"""
    try:
        data = load_grade_data(grade)
        return data
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def find_projects_for_scenario(projects_data: dict, scenario: str) -> list:
    """Find projects matching a scenario, handling different naming conventions"""
    import re
    
    if not projects_data or 'projects_by_scenario' not in projects_data:
        return []
    
    projects_by_scenario = projects_data['projects_by_scenario']
    
    # Try exact match first
    if scenario in projects_by_scenario:
        return projects_by_scenario[scenario]
    
    # Remove "Scenario X: " prefix if present
    clean_scenario = re.sub(r'^Scenario \d+:\s*', '', scenario)
    
    if clean_scenario in projects_by_scenario:
        return projects_by_scenario[clean_scenario]
    
    # Try partial matching
    for key in projects_by_scenario.keys():
        if clean_scenario in key or key in clean_scenario or key in scenario:
            return projects_by_scenario[key]
    
    return []

@api_router.get("/projects/official/{grade}/{scenario}")
async def get_official_projects(grade: str, scenario: str):
    """Get official projects for a grade and scenario"""
    try:
        projects_data = load_projects_official(grade)
        scenario_projects = find_projects_for_scenario(projects_data, scenario)
        return {"projects": scenario_projects}
    except Exception as e:
        logger.error(f"Error loading official projects: {e}")
        return {"projects": []}

@api_router.get("/projects/c21/{grade}/{scenario}")
async def get_c21_projects(grade: str, scenario: str):
    """Get C21 complementary projects for a grade and scenario"""
    try:
        projects_data = load_projects_c21(grade)
        
        if 'projects_by_scenario' in projects_data:
            # Find matching scenario (could have different naming)
            scenario_key = None
            for key in projects_data['projects_by_scenario'].keys():
                if scenario in key or key in scenario:
                    scenario_key = key
                    break
            
            if scenario_key:
                scenario_projects = projects_data['projects_by_scenario'][scenario_key]
            else:
                scenario_projects = {}
        else:
            scenario_projects = {}
            
        return {"projects": scenario_projects}
    except Exception as e:
        logger.error(f"Error loading C21 projects: {e}")
        return {"projects": {}}

@api_router.post("/planner/generate")
async def generate_planner(request: PlannerRequest):
    """Generate a complete lesson planner (Theme + Lessons) - uses pregenerated JSON if available"""
    try:
        # Load grade data for standards (always needed)
        grade_data = load_grade_data(request.grade)
        
        # Find scenario data
        scenario_data = None
        if isinstance(grade_data, list):
            scenario_data = next((s for s in grade_data if s.get('scenario', s.get('title', '')) == request.scenario), None)
        elif isinstance(grade_data, dict) and 'scenarios' in grade_data:
            scenario_data = next((s for s in grade_data['scenarios'] if s.get('scenario', s.get('title', '')) == request.scenario), None)
        
        # Extract standards and competences from curriculum
        curriculum_standards = scenario_data.get('standards_and_learning_outcomes', {}) if scenario_data else {}
        curriculum_competences = scenario_data.get('communicative_competences', {}) if scenario_data else {}
        
        # First, try to load pregenerated planner
        pregenerated = load_pregenerated_planner(
            request.grade, request.scenario, request.theme, 
            request.plan_type, request.language
        )
        
        if pregenerated:
            logger.info("Using pregenerated planner from JSON file")
            
            # Get CEFR level from grade
            cefr_levels = {
                "pre_k": "Pre A1.1",
                "K": "Pre A1.2",
                "1": "A1.1",
                "2": "A1.2",
                "3": "A2.1",
                "4": "A2.2",
                "5": "B1.1",
                "6": "B1.2"
            }
            
            # ALWAYS add standards and competences from curriculum data
            if 'theme_planner' not in pregenerated:
                pregenerated['theme_planner'] = {}
            pregenerated['theme_planner']['standards_and_learning_outcomes'] = curriculum_standards
            pregenerated['theme_planner']['communicative_competences'] = curriculum_competences
            
            # Update general_information with ALL required MEDUCA fields
            if 'general_information' not in pregenerated['theme_planner']:
                pregenerated['theme_planner']['general_information'] = {}
            
            gen_info = pregenerated['theme_planner']['general_information']
            gen_info['teachers'] = request.teacher_name or gen_info.get('teachers', '')
            gen_info['grade'] = request.grade
            gen_info['cefr_level'] = cefr_levels.get(request.grade, "A1")
            gen_info['trimester'] = request.trimester or gen_info.get('trimester', '')
            gen_info['weekly_hours'] = request.weekly_hours or gen_info.get('weekly_hours', '')
            gen_info['scenario'] = request.scenario
            gen_info['theme'] = request.theme
            
            # Build week_range
            if request.week_from and request.week_to:
                gen_info['week_range'] = f"From week {request.week_from} to week {request.week_to}"
            elif request.week_from:
                gen_info['week_range'] = f"Week {request.week_from}"
            else:
                gen_info['week_range'] = gen_info.get('week_range', '')
            
            # Add learning_sequence section if not present
            if 'learning_sequence' not in pregenerated['theme_planner']:
                project_data = None
                if request.project_id:
                    official_projects = load_projects_official(request.grade)
                    projects_list = find_projects_for_scenario(official_projects, request.scenario)
                    project_data = next((p for p in projects_list if p.get('id') == request.project_id), None)
                
                pregenerated['theme_planner']['learning_sequence'] = {
                    "project_title": project_data.get('name', '') if project_data else '',
                    "project_description": project_data.get('overview', '') if project_data else '',
                    "connection_to_theme": request.theme,
                    "lesson_dates": ["", "", "", "", ""]
                }
            
            # Add MEDUCA fields to root level for export
            pregenerated['teacher_name'] = request.teacher_name
            pregenerated['trimester'] = request.trimester
            pregenerated['weekly_hours'] = request.weekly_hours
            pregenerated['week_from'] = request.week_from
            pregenerated['week_to'] = request.week_to
            
            # Load project if specified
            if request.project_id:
                official_projects = load_projects_official(request.grade)
                projects_list = find_projects_for_scenario(official_projects, request.scenario)
                project_data = next((p for p in projects_list if p.get('id') == request.project_id), None)
                pregenerated['project'] = project_data
            
            # IMPORTANT: Update lesson_planners specific_objectives to use SMART objectives from theme_planner
            smart_objectives = pregenerated.get('theme_planner', {}).get('specific_objectives', {})
            skills_map = {1: 'listening', 2: 'reading', 3: 'speaking', 4: 'writing', 5: 'mediation'}
            
            if 'lesson_planners' in pregenerated and smart_objectives:
                for lesson in pregenerated['lesson_planners']:
                    lesson_num = lesson.get('lesson_number', 0)
                    skill_key = skills_map.get(lesson_num, '').lower()
                    if skill_key and skill_key in smart_objectives:
                        lesson['specific_objective'] = smart_objectives[skill_key]
            
            return pregenerated
        
        # If no pregenerated planner, generate basic structure from curriculum
        logger.info("No pregenerated planner found, generating basic structure from curriculum")
        
        if not scenario_data:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        # Load project if specified
        project_data = None
        if request.project_id:
            official_projects = load_projects_official(request.grade)
            projects_list = find_projects_for_scenario(official_projects, request.scenario)
            project_data = next((p for p in projects_list if p.get('id') == request.project_id), None)
        
        # Generate theme planner first to get SMART objectives
        theme_planner = generate_basic_theme_planner(scenario_data, request.theme, request.grade, request.plan_type, project_data, request)
        
        # Get SMART objectives from theme planner
        smart_objectives = theme_planner.get('specific_objectives', {})
        
        planner = {
            "grade": request.grade,
            "scenario": request.scenario,
            "theme": request.theme,
            "plan_type": request.plan_type,
            "official_format": request.official_format,
            "language": request.language,
            # Pass MEDUCA fields
            "teacher_name": request.teacher_name,
            "trimester": request.trimester,
            "weekly_hours": request.weekly_hours,
            "week_from": request.week_from,
            "week_to": request.week_to,
            "theme_planner": theme_planner,
            "lesson_planners": generate_basic_lesson_planners(scenario_data, request.theme, request.plan_type, smart_objectives),
            "project": project_data
        }
        
        return planner
        
    except Exception as e:
        logger.error(f"Error generating planner: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def generate_basic_theme_planner(scenario_data: dict, theme: str, grade: str, 
                                 plan_type: str, project_data: dict = None, request: PlannerRequest = None) -> dict:
    """Generate basic theme planner structure from curriculum data using Backward Planning approach"""
    # Get CEFR level from grade
    cefr_levels = {
        "pre_k": "Pre A1.1",
        "K": "Pre A1.2",
        "1": "A1.1",
        "2": "A1.2",
        "3": "A2.1",
        "4": "A2.2",
        "5": "B1.1",
        "6": "B1.2",
        "7": "B1.1",
        "8": "B1.2",
        "9": "B2.1",
        "10": "B2.1",
        "11": "B2.2",
        "12": "B2.2"
    }
    
    standards = scenario_data.get('standards_and_learning_outcomes', {})
    competences = scenario_data.get('communicative_competences', {})
    
    # Generate SMART objectives - format: "Students will be able to..."
    objectives = {}
    smart_templates = {
        'listening': "Students will be able to identify and understand at least 5 key vocabulary words related to {theme} when presented orally, achieving at least 80% accuracy according to a teacher observation checklist.",
        'reading': "Students will be able to read and recognize at least 5 key words related to {theme} in simple illustrated texts, achieving at least 80% accuracy according to a reading comprehension rubric.",
        'speaking': "Students will be able to orally express at least 3 complete sentences about {theme} using appropriate vocabulary and grammar structures, achieving at least 80% accuracy according to a speaking rubric.",
        'writing': "Students will be able to write at least 3 simple sentences related to {theme} with visual support, achieving at least 80% accuracy according to a writing checklist.",
        'mediation': "Students will be able to communicate key information about {theme} to peers using gestures, images, and expressions, demonstrating at least 80% comprehension according to peer feedback."
    }
    
    # Generate CEFR Can-Do Learning Outcomes - format: "Can + base verb"
    cefr_can_do_templates = {
        'listening': "Can identify key vocabulary and understand simple instructions related to {theme}.",
        'reading': "Can recognize and read simple words and short texts about {theme}.",
        'speaking': "Can express simple ideas and respond to questions about {theme}.",
        'writing': "Can write simple words and short sentences about {theme}.",
        'mediation': "Can convey basic information about {theme} using gestures and visual aids."
    }
    
    for skill in ['listening', 'reading', 'speaking', 'writing', 'mediation']:
        skill_data = standards.get(skill, {})
        if isinstance(skill_data, dict):
            learning_outcomes = skill_data.get('learning_outcomes', [])
            
            base_objective = smart_templates.get(skill, '').format(theme=theme)
            if learning_outcomes and isinstance(learning_outcomes, list) and len(learning_outcomes) > 0:
                outcome_detail = learning_outcomes[0]
                objectives[skill] = f"{base_objective} Specifically: {outcome_detail}"
            else:
                objectives[skill] = base_objective
        else:
            objectives[skill] = smart_templates.get(skill, '').format(theme=theme)
    
    # Generate materials list in ENGLISH
    materials = []
    if competences.get('linguistic'):
        ling = competences['linguistic']
        if ling.get('vocabulary'):
            vocab_items = ling['vocabulary'][:8] if isinstance(ling['vocabulary'], list) else [ling['vocabulary']]
            materials.append(f"Vocabulary flashcards: {', '.join(str(v) for v in vocab_items)}")
        if ling.get('grammatical_features') or ling.get('grammar'):
            grammar = ling.get('grammatical_features', ling.get('grammar', []))
            if isinstance(grammar, list) and grammar:
                materials.append(f"Grammar materials: {', '.join(str(g) for g in grammar[:3])}")
    
    # Add project-specific materials if available
    if project_data and project_data.get('materials'):
        proj_materials = project_data['materials']
        if isinstance(proj_materials, list):
            materials.extend(proj_materials[:5])
    
    # Add standard materials in ENGLISH
    materials.extend([
        "Whiteboard and markers",
        "Student worksheets",
        "Audio/video materials related to the theme",
        "Realia or authentic materials related to the topic",
        "Visual schedule or calendar",
        "Assessment rubrics"
    ])
    
    # Build week_range from request if available
    week_range = ""
    if request:
        if request.week_from and request.week_to:
            week_range = f"From week {request.week_from} to week {request.week_to}"
        elif request.week_from:
            week_range = f"Week {request.week_from}"
    
    theme_planner = {
        "general_information": {
            "teachers": request.teacher_name if request else "",
            "grade": grade,
            "cefr_level": cefr_levels.get(grade, "A1"),
            "trimester": request.trimester if request else "",
            "weekly_hours": request.weekly_hours if request else "",
            "week_range": week_range,
            "scenario": scenario_data.get('scenario', scenario_data.get('title', '')),
            "theme": theme
        },
        "standards_and_learning_outcomes": standards,
        "communicative_competences": competences,
        "specific_objectives": objectives,
        "materials_and_strategies": {
            "required_materials": materials,
            "differentiated_instruction": "Adapt activities for different learning styles (visual, auditory, kinesthetic). Provide scaffolding for struggling learners and extension activities for advanced students. Use pair/group work for peer support. Consider individual needs and learning paces. Offer multiple means of representation, expression, and engagement."
        },
        "learning_sequence": {
            "project_title": project_data.get('name', '') if project_data else '',
            "project_description": project_data.get('overview', '') if project_data else '',
            "connection_to_theme": theme,
            # Add lesson dates for each of the 5 lessons
            "lesson_dates": ["", "", "", "", ""]
        }
    }
    
    return theme_planner


def convert_to_can_do_format(text: str, skill: str, theme: str) -> str:
    """Convert learning outcome text to CEFR Can-Do Statement format.
    
    Rules:
    1. Must begin with "Can + base verb"
    2. No measurable time indicators
    3. No performance percentages
    4. Student-centered and skill-based
    """
    import re
    
    # Remove percentages and time indicators
    text = re.sub(r'with \d+% accuracy', '', text, flags=re.IGNORECASE)
    text = re.sub(r'achieving at least \d+% accuracy', '', text, flags=re.IGNORECASE)
    text = re.sub(r'by the end of the lesson', '', text, flags=re.IGNORECASE)
    text = re.sub(r'within \d+ minutes', '', text, flags=re.IGNORECASE)
    text = re.sub(r'according to [^.]+\.?', '', text, flags=re.IGNORECASE)
    
    # Remove teacher-centered phrasing
    text = re.sub(r'Students will be able to', '', text, flags=re.IGNORECASE)
    text = re.sub(r'The student will', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Learners will', '', text, flags=re.IGNORECASE)
    
    # Clean up the text
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)  # Remove multiple spaces
    text = re.sub(r'^\s*,\s*', '', text)  # Remove leading comma
    text = re.sub(r'\s*,\s*$', '', text)  # Remove trailing comma
    
    # If text is empty or too short, use default template
    if len(text) < 10:
        default_templates = {
            'listening': f"Can identify key vocabulary related to {theme}.",
            'reading': f"Can read and understand simple texts about {theme}.",
            'speaking': f"Can express ideas about {theme}.",
            'writing': f"Can write simple sentences about {theme}.",
            'mediation': f"Can convey information about {theme} to peers."
        }
        return default_templates.get(skill, f"Can demonstrate {skill} skills related to {theme}.")
    
    # Ensure it starts with "Can"
    if not text.lower().startswith('can '):
        # Try to extract the verb and convert to Can + verb
        text = f"Can {text[0].lower()}{text[1:]}"
    
    # Ensure first letter of "Can" is capitalized
    text = text[0].upper() + text[1:]
    
    # Ensure it ends with a period
    if not text.endswith('.'):
        text += '.'
    
    return text

def generate_basic_lesson_planners(scenario_data: dict, theme: str, plan_type: str, smart_objectives: dict = None) -> list:
    """Generate basic lesson planner structures with curriculum info"""
    lessons = []
    skills = ['listening', 'reading', 'speaking', 'writing', 'mediation']
    standards = scenario_data.get('standards_and_learning_outcomes', {})
    competences = scenario_data.get('communicative_competences', {})
    
    # CEFR Can-Do Statement templates for Learning Outcomes
    cefr_can_do_templates = {
        'listening': "Can identify key vocabulary and understand simple instructions related to {theme}.",
        'reading': "Can recognize and read simple words and short texts about {theme}.",
        'speaking': "Can express simple ideas and respond to questions about {theme}.",
        'writing': "Can write simple words and short sentences about {theme}.",
        'mediation': "Can convey basic information about {theme} using gestures and visual aids."
    }
    
    # Get vocabulary and grammar from competences
    vocab_list = []
    grammar_list = []
    if competences.get('linguistic'):
        ling = competences['linguistic']
        if isinstance(ling.get('vocabulary'), list):
            vocab_list = ling['vocabulary'][:10]  # First 10 vocab items
        elif ling.get('vocabulary'):
            vocab_list = [str(ling['vocabulary'])]
            
        if isinstance(ling.get('grammatical_features'), list):
            grammar_list = ling['grammatical_features'][:5]
        elif isinstance(ling.get('grammar'), list):
            grammar_list = ling['grammar'][:5]
    
    for i, skill in enumerate(skills, 1):
        # Get learning outcome and standards from curriculum data
        curriculum_outcome = ""
        specific_standard = ""
        if skill in standards:
            skill_data = standards[skill]
            if isinstance(skill_data, dict):
                outcomes = skill_data.get('learning_outcomes', [])
                if isinstance(outcomes, list) and outcomes:
                    curriculum_outcome = outcomes[0]
                
                # Get specific standard
                specific_standard = skill_data.get('specific', skill_data.get('general', ''))
        
        # Convert curriculum outcome to Can-Do format if it exists, otherwise use template
        if curriculum_outcome:
            # Convert to Can-Do format: "Can + verb..."
            learning_outcome = convert_to_can_do_format(curriculum_outcome, skill, theme)
        else:
            # Use default Can-Do template
            learning_outcome = cefr_can_do_templates.get(skill, '').format(theme=theme)
        
        # Generate basic activities based on curriculum content
        warm_up_activities = [
            f"Review {theme} vocabulary: {', '.join(vocab_list[:5]) if vocab_list else 'key terms'}",
            f"Activate prior knowledge about {theme}",
            "Engage students with visuals or realia related to the topic"
        ]
        
        presentation_activities = [
            f"Introduce {skill} focus using authentic materials",
            f"Present key vocabulary: {', '.join(vocab_list[:5]) if vocab_list else 'topic-related words'}",
            f"Model {skill} skill with clear examples"
        ]
        
        if grammar_list:
            presentation_activities.append(f"Introduce grammar: {', '.join(grammar_list[:2])}")
        
        preparation_activities = [
            f"Guided {skill} practice with scaffolding",
            "Work in pairs or small groups",
            f"Practice using {theme} vocabulary in context"
        ]
        
        performance_activities = [
            f"Independent {skill} activity",
            f"Students demonstrate {skill} skill with {theme} content",
            "Provide opportunities for student production"
        ]
        
        assessment_activities = [
            f"Formative assessment of {skill} skill",
            "Check understanding and provide feedback",
            "Students self-assess their performance"
        ]
        
        reflection_activities = [
            f"Reflect on {skill} learning",
            f"Discuss what was learned about {theme}",
            "Set goals for improvement"
        ]
        
        # Get SMART objective from theme planner if available
        smart_objective = ""
        if smart_objectives and skill in smart_objectives:
            smart_objective = smart_objectives[skill]
        
        lesson = {
            "lesson_number": i,
            "skill_focus": skill.capitalize(),
            "scenario": scenario_data.get('scenario', scenario_data.get('title', '')),
            "theme": theme,
            "date": "",
            "time": "45-60 minutes",
            "specific_objective": smart_objective if smart_objective else (specific_standard if specific_standard else f"Students will develop {skill} skills related to {theme}"),
            "learning_outcome": learning_outcome,  # Already in Can-Do format from convert_to_can_do_format()
            "lesson_stages": [
                {
                    "stage": "Warm-up / Pre-task",
                    "activities": warm_up_activities
                },
                {
                    "stage": "Presentation",
                    "activities": presentation_activities
                },
                {
                    "stage": "Preparation",
                    "activities": preparation_activities
                },
                {
                    "stage": "Performance",
                    "activities": performance_activities
                },
                {
                    "stage": "Assessment / Post-task",
                    "activities": assessment_activities
                },
                {
                    "stage": "Reflection",
                    "activities": reflection_activities
                }
            ],
            "comments": {
                "homework": f"Practice {skill} skill with {theme} content at home. Review vocabulary and grammar.",
                "formative_assessment": f"Observe student {skill} performance. Check comprehension and production of {theme} content.",
                "teacher_comments": f"Note: This is a basic structure from curriculum. Complete with specific activities for {skill} and {theme}. Consider student level and needs."
            }
        }
        lessons.append(lesson)
    
    return lessons

async def generate_theme_planner_with_ai(scenario_data: dict, theme: str, grade: str, 
                                         plan_type: str, language: str, project_data: dict = None) -> dict:
    """Generate theme planner with AI-generated content"""
    # Get CEFR level from grade
    cefr_levels = {
        "pre_k": "Pre A1.1",
        "K": "Pre A1.2",
        "1": "A1.1",
        "2": "A1.2",
        "3": "A2.1",
        "4": "A2.2",
        "5": "B1.1",
        "6": "B1.2"
    }
    
    standards = scenario_data.get('standards_and_learning_outcomes', {})
    competences = scenario_data.get('communicative_competences', {})
    
    # Create AI prompt for SMART objectives
    objectives_prompt = create_objectives_prompt(scenario_data, theme, grade, language, standards)
    
    # Generate SMART objectives with AI
    llm_key = os.environ.get('EMERGENT_LLM_KEY')
    chat = LlmChat(
        api_key=llm_key,
        session_id=f"objectives-{grade}-{theme}",
        system_message="You are an expert English teacher and curriculum specialist for Panama's MEDUCA curriculum."
    ).with_model("openai", "gpt-5.2")
    
    objectives_message = UserMessage(text=objectives_prompt)
    objectives_response = await chat.send_message(objectives_message)
    
    # Parse objectives response
    smart_objectives = parse_objectives_response(objectives_response)
    
    # Generate materials and strategies
    materials_prompt = create_materials_prompt(scenario_data, theme, grade, language, plan_type, project_data)
    materials_message = UserMessage(text=materials_prompt)
    materials_response = await chat.send_message(materials_message)
    materials_data = parse_materials_response(materials_response)
    
    theme_planner = {
        "general_information": {
            "teachers": "",
            "grade": grade,
            "cefr_level": cefr_levels.get(grade, "A1"),
            "trimester": "",
            "weekly_hours": "",
            "week_range": "",
            "scenario": scenario_data.get('scenario', scenario_data.get('title', '')),
            "theme": theme
        },
        "standards_and_learning_outcomes": standards,
        "communicative_competences": competences,
        "specific_objectives": smart_objectives,
        "materials_and_strategies": materials_data,
        "learning_sequence": {
            "project_title": project_data.get('name', '') if project_data else '',
            "project_description": project_data.get('overview', '') if project_data else '',
            "connection_to_theme": theme
        }
    }
    
    return theme_planner

async def generate_lesson_planners_with_ai(scenario_data: dict, theme: str, plan_type: str, 
                                           language: str, grade: str) -> list:
    """Generate 5 lesson planners with complete AI-generated content"""
    lessons = []
    skills = ['listening', 'reading', 'speaking', 'writing', 'mediation']
    standards = scenario_data.get('standards_and_learning_outcomes', {})
    competences = scenario_data.get('communicative_competences', {})
    
    llm_key = os.environ.get('EMERGENT_LLM_KEY')
    
    for i, skill in enumerate(skills, 1):
        # Create detailed prompt for this lesson
        lesson_prompt = create_lesson_prompt(
            scenario_data, theme, skill, grade, plan_type, language, 
            standards, competences, i
        )
        
        # Generate lesson content with AI
        chat = LlmChat(
            api_key=llm_key,
            session_id=f"lesson-{grade}-{theme}-{skill}",
            system_message="You are an expert English teacher creating detailed, ready-to-teach lesson plans for Panama's MEDUCA curriculum."
        ).with_model("openai", "gpt-5.2")
        
        lesson_message = UserMessage(text=lesson_prompt)
        lesson_response = await chat.send_message(lesson_message)
        
        # Parse AI response into structured lesson
        lesson_data = parse_lesson_response(lesson_response, i, skill, scenario_data, theme)
        lessons.append(lesson_data)
    
    return lessons

def create_objectives_prompt(scenario_data: dict, theme: str, grade: str, language: str, standards: dict) -> str:
    """Create prompt for generating SMART objectives"""
    scenario_title = scenario_data.get('scenario', scenario_data.get('title', ''))
    
    prompt = f"""Generate 5 SMART objectives (one for each skill: Listening, Reading, Speaking, Writing, Mediation) for this English lesson:

Grade: {grade}
Scenario: {scenario_title}
Theme: {theme}
Language: {'Spanish' if language == 'es' else 'English'}

Standards and Learning Outcomes:
{json.dumps(standards, indent=2)}

Each objective must be:
- Specific: clearly defined
- Measurable: observable outcome
- Achievable: appropriate for grade level
- Relevant: aligned to curriculum standards
- Time-bound: for this lesson/theme

Format your response as JSON:
{{
  "listening": "objective text",
  "reading": "objective text",
  "speaking": "objective text",
  "writing": "objective text",
  "mediation": "objective text"
}}

Write in {'Spanish' if language == 'es' else 'English'}."""
    
    return prompt

def create_materials_prompt(scenario_data: dict, theme: str, grade: str, language: str, 
                            plan_type: str, project_data: dict = None) -> str:
    """Create prompt for generating materials and strategies"""
    scenario_title = scenario_data.get('scenario', scenario_data.get('title', ''))
    competences = scenario_data.get('communicative_competences', {})
    
    prompt = f"""Generate a comprehensive list of materials and differentiated instruction strategies for this English lesson:

Grade: {grade}
Scenario: {scenario_title}
Theme: {theme}
Plan Type: {plan_type}
Language: {'Spanish' if language == 'es' else 'English'}

Communicative Competences:
{json.dumps(competences, indent=2, ensure_ascii=False)}

{"Project: " + project_data.get('name', '') if project_data else ""}

Provide:
1. Required Materials: specific items needed (flashcards, posters, manipulatives, tech, etc.)
2. Differentiated Instruction: strategies for diverse learners (visual, auditory, kinesthetic, struggling, advanced, ELL)

Format as JSON:
{{
  "required_materials": ["material 1", "material 2", ...],
  "differentiated_instruction": "detailed strategies text"
}}

Write in {'Spanish' if language == 'es' else 'English'}."""
    
    return prompt

def create_lesson_prompt(scenario_data: dict, theme: str, skill: str, grade: str, 
                        plan_type: str, language: str, standards: dict, 
                        competences: dict, lesson_number: int) -> str:
    """Create comprehensive prompt for generating a complete lesson"""
    scenario_title = scenario_data.get('scenario', scenario_data.get('title', ''))
    
    skill_data = standards.get(skill, {})
    learning_outcome = ""
    if isinstance(skill_data, dict):
        outcomes = skill_data.get('learning_outcomes', [])
        if isinstance(outcomes, list) and outcomes:
            learning_outcome = outcomes[0]
    
    # Get vocabulary and grammar from competences
    vocab = ""
    grammar = ""
    if competences.get('linguistic'):
        ling = competences['linguistic']
        if isinstance(ling.get('vocabulary'), list):
            vocab = ", ".join(ling['vocabulary'][:15])
        elif ling.get('vocabulary'):
            vocab = str(ling['vocabulary'])
            
        if isinstance(ling.get('grammatical_features'), list):
            grammar = ", ".join(ling['grammatical_features'])
        elif isinstance(ling.get('grammar'), list):
            grammar = ", ".join(ling['grammar'])
    
    detail_level = {
        'basic': 'concise and essential',
        'standard': 'complete and balanced',
        'enriched': 'detailed with extension activities'
    }
    
    prompt = f"""Create a complete, ready-to-teach lesson plan for:

Lesson {lesson_number}: {skill.upper()} Skill
Grade: {grade}
Scenario: {scenario_title}
Theme: {theme}
Plan Type: {plan_type} ({detail_level.get(plan_type, 'standard')})
Language: {'Spanish' if language == 'es' else 'English'}

Learning Outcome: {learning_outcome}
Key Vocabulary: {vocab}
Grammar: {grammar}

Generate {detail_level.get(plan_type, 'standard')} content for the 6 Action-Oriented Approach stages:

1. WARM-UP / PRE-TASK (5-10 min):
   - Specific engagement activity
   - How to activate prior knowledge
   - Materials needed

2. PRESENTATION (10-15 min):
   - How to introduce new content
   - Examples to use
   - Visual aids/demonstrations

3. PREPARATION (10-15 min):
   - Guided practice activities
   - Scaffolding techniques
   - Teacher support strategies

4. PERFORMANCE (15-20 min):
   - Independent practice tasks
   - Student production activities
   - Expected outcomes

5. ASSESSMENT / POST-TASK (10 min):
   - Formative assessment strategies
   - How to check understanding
   - Feedback techniques

6. REFLECTION (5 min):
   - Reflection questions
   - Self-assessment activities
   - Closure strategy

Also include:
- SPECIFIC OBJECTIVE: SMART objective for this lesson
- HOMEWORK: meaningful assignment
- TEACHER NOTES: important tips and common challenges

Format as JSON:
{{
  "specific_objective": "SMART objective",
  "lesson_stages": [
    {{
      "stage": "Warm-up / Pre-task",
      "activities": ["activity 1 with details", "activity 2 with details", ...]
    }},
    ... (6 stages total)
  ],
  "homework": "homework description",
  "formative_assessment": "assessment strategies",
  "teacher_comments": "tips and notes for teacher"
}}

Make activities specific, practical, and ready to implement. Include exact instructions the teacher can follow.
Write in {'Spanish' if language == 'es' else 'English'}."""
    
    return prompt

def parse_objectives_response(response: str) -> dict:
    """Parse AI response for objectives"""
    try:
        # Try to extract JSON from response
        import re
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
    except:
        pass
    
    # Fallback
    return {
        "listening": response[:200] if len(response) > 200 else response,
        "reading": "",
        "speaking": "",
        "writing": "",
        "mediation": ""
    }

def parse_materials_response(response: str) -> dict:
    """Parse AI response for materials"""
    try:
        import re
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group())
            return {
                "required_materials": data.get('required_materials', []),
                "differentiated_instruction": data.get('differentiated_instruction', '')
            }
    except:
        pass
    
    return {
        "required_materials": ["Materials to be determined"],
        "differentiated_instruction": response[:500] if len(response) > 500 else response
    }

def parse_lesson_response(response: str, lesson_number: int, skill: str, 
                         scenario_data: dict, theme: str) -> dict:
    """Parse AI response into structured lesson"""
    try:
        import re
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group())
            
            return {
                "lesson_number": lesson_number,
                "skill_focus": skill.capitalize(),
                "scenario": scenario_data.get('scenario', scenario_data.get('title', '')),
                "theme": theme,
                "date": "",
                "time": "45-60 minutes",
                "specific_objective": data.get('specific_objective', ''),
                "learning_outcome": data.get('learning_outcome', get_learning_outcome_fallback(scenario_data, skill)),
                "lesson_stages": data.get('lesson_stages', []),
                "comments": {
                    "homework": data.get('homework', ''),
                    "formative_assessment": data.get('formative_assessment', ''),
                    "teacher_comments": data.get('teacher_comments', '')
                }
            }
    except Exception as e:
        logger.error(f"Error parsing lesson response: {e}")
    
    # Fallback structure
    return {
        "lesson_number": lesson_number,
        "skill_focus": skill.capitalize(),
        "scenario": scenario_data.get('scenario', scenario_data.get('title', '')),
        "theme": theme,
        "date": "",
        "time": "45-60 minutes",
        "specific_objective": "To be completed",
        "learning_outcome": get_learning_outcome_fallback(scenario_data, skill),
        "lesson_stages": generate_lesson_stages('standard'),
        "comments": {
            "homework": "",
            "formative_assessment": "",
            "teacher_comments": ""
        }
    }

def get_learning_outcome_fallback(scenario_data: dict, skill: str) -> str:
    """Get learning outcome from standards"""
    standards = scenario_data.get('standards_and_learning_outcomes', {})
    if skill in standards:
        skill_data = standards[skill]
        if isinstance(skill_data, dict):
            outcomes = skill_data.get('learning_outcomes', [])
            if isinstance(outcomes, list) and outcomes:
                return outcomes[0]
    return ""

def generate_lesson_stages(plan_type: str) -> list:
    """Generate the 6 action-oriented approach stages"""
    stages = [
        {
            "stage": "Warm-up / Pre-task",
            "activities": [] if plan_type == "basic" else ["Engagement activity", "Modeling", "Clarification"]
        },
        {
            "stage": "Presentation",
            "activities": [] if plan_type == "basic" else ["Introduce new content", "Examples and demonstration"]
        },
        {
            "stage": "Preparation",
            "activities": [] if plan_type == "basic" else ["Guided practice", "Scaffolding"]
        },
        {
            "stage": "Performance",
            "activities": [] if plan_type == "basic" else ["Independent practice", "Student production"]
        },
        {
            "stage": "Assessment / Post-task",
            "activities": [] if plan_type == "basic" else ["Formative assessment", "Feedback"]
        },
        {
            "stage": "Reflection",
            "activities": [] if plan_type == "basic" else ["Student reflection", "Lesson review"]
        }
    ]
    return stages

@api_router.post("/planner/export/docx")
async def export_to_docx(planner: dict):
    """Export planner to DOCX format - MEDUCA Official Format"""
    try:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # Create a new Document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Get theme planner data
        theme_planner = planner.get('theme_planner', {})
        general_info = theme_planner.get('general_information', {})
        
        # ========================================
        # MEDUCA HEADER
        # ========================================
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = header_para.add_run("GOBIERNO NACIONAL • CON PASO FIRME • | ")
        run1.font.size = Pt(8)
        run2 = header_para.add_run("MINISTERIO DE EDUCACIÓN")
        run2.font.size = Pt(8)
        run2.bold = True
        
        subheader = doc.add_paragraph()
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subheader.add_run("Dirección Nacional de Currículo de Lengua Extranjera")
        sub_run.font.size = Pt(8)
        
        # ========================================
        # THEME PLANNER TITLE
        # ========================================
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Theme Planner # ___ – Overview")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        doc.add_paragraph()
        
        # ========================================
        # SECTION 1: GENERAL INFORMATION
        # ========================================
        doc.add_paragraph().add_run("1. General Information:").bold = True
        
        gen_table = doc.add_table(rows=5, cols=4)
        gen_table.style = 'Table Grid'
        
        # Row 1: Teachers
        gen_table.cell(0, 0).text = "1. Teacher(s):"
        gen_table.cell(0, 1).merge(gen_table.cell(0, 3)).text = general_info.get('teachers', '_______________')
        
        # Row 2: Grade, CEFR, Trimester
        gen_table.cell(1, 0).text = "2. Grade:"
        gen_table.cell(1, 1).text = general_info.get('grade', '______')
        gen_table.cell(1, 2).text = "3. CEFR Level:"
        gen_table.cell(1, 3).text = general_info.get('cefr_level', '______')
        
        # Row 3: Trimester, Weekly Hours
        gen_table.cell(2, 0).text = "4. Trimester:"
        gen_table.cell(2, 1).text = general_info.get('trimester', '______')
        gen_table.cell(2, 2).text = "5. Weekly Hour(s):"
        gen_table.cell(2, 3).text = general_info.get('weekly_hours', '______')
        
        # Row 4: Weeks
        gen_table.cell(3, 0).text = "6. Week(s):"
        gen_table.cell(3, 1).merge(gen_table.cell(3, 3)).text = general_info.get('week_range', 'From week ____ to week ____')
        
        # Row 5: Scenario, Theme
        gen_table.cell(4, 0).text = "7. Scenario __:"
        gen_table.cell(4, 1).text = general_info.get('scenario', '______')
        gen_table.cell(4, 2).text = "8. Theme __:"
        gen_table.cell(4, 3).text = general_info.get('theme', '______')
        
        # Make first column bold
        for row in gen_table.rows:
            row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
        
        doc.add_paragraph()
        
        # ========================================
        # SECTION 2: SPECIFIC STANDARDS AND LEARNING OUTCOMES
        # ========================================
        doc.add_paragraph().add_run("2. Specific Standards and Learning Outcomes:").bold = True
        
        standards_data = theme_planner.get('standards_and_learning_outcomes', {})
        
        std_table = doc.add_table(rows=6, cols=3)
        std_table.style = 'Table Grid'
        
        # Header row
        std_table.cell(0, 0).text = "Skills:"
        std_table.cell(0, 1).text = "Specific Standards:"
        std_table.cell(0, 2).text = "Learning Outcomes:"
        for cell in std_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
        
        skills_list = [
            ('1. Listening:', 'listening'),
            ('2. Reading:', 'reading'),
            ('3. Speaking:', 'speaking'),
            ('4. Writing:', 'writing'),
            ('5. Mediation:', 'mediation')
        ]
        
        for idx, (label, key) in enumerate(skills_list, 1):
            skill_data = standards_data.get(key, {})
            standards = []
            outcomes = []
            
            if isinstance(skill_data, dict):
                # Get learning_outcomes if present (Grade 1+ format)
                outcomes = skill_data.get('learning_outcomes', [])
                
                # Get standards - check both formats
                if isinstance(skill_data.get('specific'), list):
                    # Grade 1+ format
                    standards = skill_data['specific']
                else:
                    # K/Pre-K format: individual fields
                    fields = ['receptive', 'interactive', 'productive', 'reading', 'reading1', 'reading2', 
                             'phonemic_awareness', 'listening', 'listening1', 'listening2', 'speaking', 'speaking1', 
                             'speaking2', 'writing', 'writing1', 'writing2', 'text', 'concept', 'general']
                    for field in fields:
                        val = skill_data.get(field)
                        if val and isinstance(val, str):
                            standards.append(val)
                
                # If no learning_outcomes, generate from standards in Can-Do format (K/Pre-K format)
                if not outcomes and standards:
                    outcomes = [convert_to_can_do_format(s, key, theme_planner.get('general_information', {}).get('theme', '')) for s in standards]
            
            std_table.cell(idx, 0).text = label
            std_table.cell(idx, 0).paragraphs[0].runs[0].bold = True
            std_table.cell(idx, 1).text = '\n• '.join([''] + standards) if standards else 'To be defined'
            std_table.cell(idx, 2).text = '\n• '.join([''] + outcomes) if outcomes else 'To be defined'
        
        doc.add_paragraph()
        
        # ========================================
        # SECTION 3: COMMUNICATIVE COMPETENCES
        # ========================================
        doc.add_paragraph().add_run("3. Communicative Competences").bold = True
        
        comp_table = doc.add_table(rows=2, cols=3)
        comp_table.style = 'Table Grid'
        
        # Header row
        comp_table.cell(0, 0).text = "Linguistic Competence\n(Learn to Know)"
        comp_table.cell(0, 1).text = "Pragmatic Competence\n(Learn to Do)"
        comp_table.cell(0, 2).text = "Sociolinguistic Competence\n(Learn to Be)"
        for cell in comp_table.rows[0].cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
        
        # Content row
        competences = theme_planner.get('communicative_competences', {})
        ling = competences.get('linguistic', {})
        
        # Linguistic content
        ling_content = ""
        if isinstance(ling, dict):
            grammar = ling.get('grammatical_features', ling.get('grammar', []))
            if isinstance(grammar, list):
                grammar = ', '.join(grammar)
            ling_content += f"• Grammatical Features:\n{grammar}\n\n"
            
            vocab = ling.get('vocabulary', [])
            if isinstance(vocab, list):
                vocab = ', '.join(vocab[:10])
            ling_content += f"• Vocabulary:\n{vocab}\n\n"
            
            phonemic = ling.get('phonemic_awareness', ling.get('pronunciation', ''))
            ling_content += f"• Pronunciation & Phonemic Awareness:\n{phonemic}"
        
        comp_table.cell(1, 0).text = ling_content
        
        # Pragmatic content
        prag = competences.get('pragmatic', '')
        if isinstance(prag, list):
            prag = ', '.join(prag)
        elif isinstance(prag, dict):
            prag = prag.get('functions', str(prag))
        comp_table.cell(1, 1).text = str(prag) if prag else '______'
        
        # Sociolinguistic content
        socio = competences.get('sociolinguistic', '')
        if isinstance(socio, list):
            socio = ', '.join(socio)
        comp_table.cell(1, 2).text = str(socio) if socio else '______'
        
        # 21st Century Project
        p = doc.add_paragraph()
        p.add_run("\n21st-Century Skills Project").bold = True
        
        project_data = planner.get('project')
        if project_data:
            doc.add_paragraph(f"Name: {project_data.get('name', '')}")
            doc.add_paragraph(f"Category: {project_data.get('category', '')}")
            doc.add_paragraph(f"Overview: {project_data.get('overview', '')}")
        else:
            doc.add_paragraph("No project selected")
        
        doc.add_paragraph()
        
        # ========================================
        # SECTION 4: SPECIFIC OBJECTIVES
        # ========================================
        doc.add_paragraph().add_run("4. Specific Objectives").bold = True
        
        obj_table = doc.add_table(rows=5, cols=2)
        obj_table.style = 'Table Grid'
        
        objectives = theme_planner.get('specific_objectives', {})
        for idx, skill in enumerate(['Listening', 'Reading', 'Speaking', 'Writing', 'Mediation']):
            obj_table.cell(idx, 0).text = skill
            obj_table.cell(idx, 0).paragraphs[0].runs[0].bold = True
            obj_table.cell(idx, 1).text = objectives.get(skill.lower(), 'To be completed')
        
        doc.add_paragraph()
        
        # ========================================
        # SECTION 5: MATERIALS AND TEACHING STRATEGIES
        # ========================================
        doc.add_paragraph().add_run("5. Materials and Teaching Strategies").bold = True
        
        mat_table = doc.add_table(rows=2, cols=2)
        mat_table.style = 'Table Grid'
        
        materials_data = theme_planner.get('materials_and_strategies', {})
        
        mat_table.cell(0, 0).text = "Materials"
        mat_table.cell(0, 0).paragraphs[0].runs[0].bold = True
        materials_list = materials_data.get('required_materials', [])
        mat_table.cell(0, 1).text = '\n• '.join([''] + materials_list) if materials_list else 'To be completed'
        
        mat_table.cell(1, 0).text = "Differentiated Instruction and\nAccommodations for Students with\nDiverse Learning Needs (DLN)"
        mat_table.cell(1, 0).paragraphs[0].runs[0].bold = True
        mat_table.cell(1, 1).text = materials_data.get('differentiated_instruction', 'To be completed')
        
        doc.add_paragraph()
        
        # ========================================
        # SECTION 6: LEARNING SEQUENCE
        # ========================================
        doc.add_paragraph().add_run("6. Learning Sequence").bold = True
        
        lesson_planners = planner.get('lesson_planners', [])
        num_lessons = len(lesson_planners) if lesson_planners else 5
        
        seq_table = doc.add_table(rows=num_lessons + 1, cols=3)
        seq_table.style = 'Table Grid'
        
        # Header
        seq_table.cell(0, 0).text = "Lesson"
        seq_table.cell(0, 1).text = "Learning Sequence"
        seq_table.cell(0, 2).text = "Date"
        for cell in seq_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        lesson_dates = theme_planner.get('learning_sequence', {}).get('lesson_dates', ['', '', '', '', ''])
        
        # Use specific objectives from lesson planners as the learning sequence description
        for idx, lesson in enumerate(lesson_planners):
            skill = lesson.get('skill_focus', '')
            # Use the specific objective as the learning sequence description
            description = f"{skill}: {lesson.get('specific_objective', 'To be defined')}"
            
            seq_table.cell(idx + 1, 0).text = f"Lesson {lesson.get('lesson_number', idx + 1)}"
            seq_table.cell(idx + 1, 0).paragraphs[0].runs[0].bold = True
            seq_table.cell(idx + 1, 1).text = description
            seq_table.cell(idx + 1, 2).text = lesson_dates[idx] if idx < len(lesson_dates) else '______'
        
        # ========================================
        # LESSON PLANNERS - Page break and new section
        # ========================================
        for lesson in planner.get('lesson_planners', []):
            doc.add_page_break()
            
            # MEDUCA Header for Lesson Planner
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h_run = header.add_run("GOBIERNO NACIONAL • CON PASO FIRME • | MINISTERIO DE EDUCACIÓN")
            h_run.font.size = Pt(8)
            
            # Title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t_run = title.add_run(f"Lesson Planner – Theme # ___ – Lesson # {lesson['lesson_number']}")
            t_run.bold = True
            t_run.font.size = Pt(14)
            t_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Instructions
            inst = doc.add_paragraph()
            inst.add_run("Instructions: ").bold = True
            inst.add_run("Complete this planner five times per theme, once per lesson (Lesson 1, Lesson 2, Lesson 3, Lesson 4, and Lesson 5).")
            inst.runs[-1].font.size = Pt(9)
            inst.runs[-1].italic = True
            
            # Lesson header info
            p = doc.add_paragraph()
            p.add_run(f"Lesson # {lesson['lesson_number']}").bold = True
            p.add_run(f"     Skills focus for this lesson: {lesson['skill_focus']}")
            
            # Info table
            info_table = doc.add_table(rows=4, cols=6)
            info_table.style = 'Table Grid'
            
            # Row 1: Grade, Scenario, Theme
            info_table.cell(0, 0).text = "Grade:"
            info_table.cell(0, 0).paragraphs[0].runs[0].bold = True
            info_table.cell(0, 1).text = planner.get('grade', '')
            info_table.cell(0, 2).text = "Scenario __:"
            info_table.cell(0, 2).paragraphs[0].runs[0].bold = True
            info_table.cell(0, 3).text = lesson.get('scenario', '')
            info_table.cell(0, 4).text = "Theme __:"
            info_table.cell(0, 4).paragraphs[0].runs[0].bold = True
            info_table.cell(0, 5).text = lesson.get('theme', '')
            
            # Row 2: Date, Learning Sequence time
            info_table.cell(1, 0).text = "Date(s):"
            info_table.cell(1, 0).paragraphs[0].runs[0].bold = True
            info_table.cell(1, 1).merge(info_table.cell(1, 2)).text = lesson.get('date', '______')
            info_table.cell(1, 3).text = "Learning Sequence time:"
            info_table.cell(1, 3).paragraphs[0].runs[0].bold = True
            info_table.cell(1, 4).merge(info_table.cell(1, 5)).text = lesson.get('time', '45-60 minutes')
            
            # Row 3: Specific Objective
            info_table.cell(2, 0).text = "Specific Objective:"
            info_table.cell(2, 0).paragraphs[0].runs[0].bold = True
            info_table.cell(2, 1).merge(info_table.cell(2, 5)).text = lesson.get('specific_objective', '______')
            
            # Row 4: Learning Outcome
            info_table.cell(3, 0).text = "Learning Outcome:"
            info_table.cell(3, 0).paragraphs[0].runs[0].bold = True
            info_table.cell(3, 1).merge(info_table.cell(3, 5)).text = lesson.get('learning_outcome', '______')
            
            doc.add_paragraph()
            
            # Six Action-oriented Approach Lesson Stages
            stages_table = doc.add_table(rows=7, cols=2)
            stages_table.style = 'Table Grid'
            
            # Header
            stages_table.cell(0, 0).text = "The Six Action-oriented Approach Lesson Stages"
            stages_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            stages_table.cell(0, 0).paragraphs[0].runs[0].bold = True
            stages_table.cell(0, 1).text = "Estimated\nDate and Time"
            stages_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            stages_table.cell(0, 1).paragraphs[0].runs[0].bold = True
            
            stage_names = [
                "Stage 1 - Warm-up / Pre-task (Engagement, Modeling and Clarification):",
                "Stage 2 - Presentation:",
                "Stage 3 - Preparation:",
                "Stage 4 - Performance:",
                "Stage 5 - Assessment / Post-task:",
                "Stage 6 - Reflection:"
            ]
            
            for idx, stage_name in enumerate(stage_names, 1):
                stage_data = lesson.get('lesson_stages', [])[idx - 1] if idx <= len(lesson.get('lesson_stages', [])) else {}
                activities = stage_data.get('activities', [])
                
                cell_content = f"{stage_name}\n"
                if activities:
                    cell_content += '\n'.join([f"• {act}" for act in activities])
                else:
                    cell_content += "• To be completed"
                
                stages_table.cell(idx, 0).text = cell_content
                stages_table.cell(idx, 1).text = stage_data.get('estimated_time', '______')
                stages_table.cell(idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Comments and Observations
            comments_table = doc.add_table(rows=4, cols=2)
            comments_table.style = 'Table Grid'
            
            comments_table.cell(0, 0).merge(comments_table.cell(0, 1)).text = "Comments and Observations"
            comments_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            comments_table.cell(0, 0).paragraphs[0].runs[0].bold = True
            
            comments = lesson.get('comments', {})
            
            comments_table.cell(1, 0).text = "Homework:"
            comments_table.cell(1, 0).paragraphs[0].runs[0].bold = True
            comments_table.cell(1, 1).text = comments.get('homework', '______')
            
            comments_table.cell(2, 0).text = "Formative Assessment of Learning:"
            comments_table.cell(2, 0).paragraphs[0].runs[0].bold = True
            comments_table.cell(2, 1).text = comments.get('formative_assessment', '______')
            
            comments_table.cell(3, 0).text = "Teacher's Comments/Observations:"
            comments_table.cell(3, 0).paragraphs[0].runs[0].bold = True
            comments_table.cell(3, 1).text = comments.get('teacher_comments', '______')
        
        # ========================================
        # 21ST CENTURY PROJECT - Full details if exists
        # ========================================
        if project_data:
            doc.add_page_break()
            doc.add_heading('21st Century Project (Lesson 5) - Full Details', 1)
            
            p = doc.add_paragraph()
            p.add_run('Project Name: ').bold = True
            p.add_run(project_data.get('name', ''))
            
            p = doc.add_paragraph()
            p.add_run('Category: ').bold = True
            p.add_run(project_data.get('category', ''))
            
            p = doc.add_paragraph()
            p.add_run('Overview: ').bold = True
            p.add_run(project_data.get('overview', ''))
            
            p = doc.add_paragraph()
            p.add_run('General Objective: ').bold = True
            p.add_run(project_data.get('general_objective', ''))
            
            specific_objs = project_data.get('specific_objectives', [])
            if specific_objs:
                doc.add_heading('Specific Objectives', 3)
                for obj in specific_objs:
                    doc.add_paragraph(f"• {obj}", style='List Bullet')
            
            activities = project_data.get('activities', [])
            if activities:
                doc.add_heading('Project Activities', 3)
                for act in activities:
                    doc.add_paragraph(f"• {act}", style='List Bullet')
            
            products = project_data.get('products_evidences', [])
            if products:
                doc.add_heading('Products / Evidences', 3)
                for prod in products:
                    doc.add_paragraph(f"• {prod}", style='List Bullet')
            
            rubric = project_data.get('rubric_criteria', {})
            if rubric:
                doc.add_heading('Rubric Criteria', 3)
                for criterion, description in rubric.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{criterion}: ").bold = True
                    p.add_run(str(description))
            
            p = doc.add_paragraph()
            p.add_run('Duration: ').bold = True
            p.add_run(project_data.get('duration', ''))
            
            skills = project_data.get('skills_developed', [])
            if skills:
                p = doc.add_paragraph()
                p.add_run('Skills Developed: ').bold = True
                p.add_run(', '.join(skills))
        
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Save to temp file and return
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(file_stream.getvalue())
            tmp_path = tmp.name
        
        return FileResponse(
            tmp_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"lesson_planner_{planner['grade']}_{planner['scenario']}.docx"
        )
        
    except Exception as e:
        logger.error(f"Error exporting to DOCX: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
