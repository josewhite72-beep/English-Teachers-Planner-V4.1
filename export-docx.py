"""
Vercel Serverless Function - Export Planner to Word
Genera documentos .docx usando python-docx
"""

from http.server import BaseHTTPRequestHandler
import json
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            # Leer datos del request
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            planner_data = json.loads(post_data.decode('utf-8'))
            
            # Generar documento
            doc = create_planner_document(planner_data)
            
            # Guardar en memoria
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            # Enviar respuesta
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="lesson_planner_{planner_data.get("grade", "unknown")}.docx"')
            self.end_headers()
            self.wfile.write(doc_io.getvalue())
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            error_response = json.dumps({'error': str(e)})
            self.wfile.write(error_response.encode())


def create_planner_document(data):
    """Crear documento Word completo"""
    doc = Document()
    
    # Header MEDUCA
    add_meduca_header(doc)
    
    # Theme Planner
    add_theme_planner(doc, data)
    
    # Lesson Planners
    add_lesson_planners(doc, data)
    
    return doc


def add_meduca_header(doc):
    """Header oficial MEDUCA"""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run1 = heading.add_run('MINISTERIO DE EDUCACIÓN\n')
    run1.bold = True
    run1.font.size = Pt(14)
    
    run2 = heading.add_run('Dirección Nacional de Currículo de Lengua Extranjera')
    run2.font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio


def add_theme_planner(doc, data):
    """Theme Planner completo"""
    theme = data.get('theme_planner', {})
    
    # Título
    title = doc.add_heading('Theme Planner - Overview', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. General Information
    doc.add_heading('1. General Information:', level=2)
    add_general_info_table(doc, theme.get('general_information', {}))
    
    # 2. Standards and Learning Outcomes
    doc.add_heading('2. Specific Standards and Learning Outcomes:', level=2)
    add_standards_table(doc, theme.get('standards_and_learning_outcomes', {}))
    
    # 3. Communicative Competences
    doc.add_heading('3. Communicative Competences', level=2)
    add_competences_table(doc, theme.get('communicative_competences', {}))
    
    # Project
    if 'project' in data:
        add_project_section(doc, data['project'])
    
    # 4. Specific Objectives
    doc.add_heading('4. Specific Objectives', level=2)
    add_objectives_table(doc, theme.get('specific_objectives', {}))
    
    # 5. Materials
    doc.add_heading('5. Materials and Teaching Strategies', level=2)
    add_materials_table(doc, theme.get('materials_and_strategies', {}))
    
    # 6. Learning Sequence
    doc.add_heading('6. Learning Sequence', level=2)
    add_learning_sequence_table(doc, data.get('lesson_planners', []), theme.get('learning_sequence', {}))


def add_general_info_table(doc, general):
    """Tabla de información general"""
    grade_labels = {
        'pre_k': 'Pre-K', 'K': 'Kindergarten',
        '1': 'Grade 1', '2': 'Grade 2', '3': 'Grade 3', '4': 'Grade 4',
        '5': 'Grade 5', '6': 'Grade 6', '7': 'Grade 7', '8': 'Grade 8',
        '9': 'Grade 9', '10': 'Grade 10', '11': 'Grade 11', '12': 'Grade 12',
    }
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    # Row 1: Teacher
    table.cell(0, 0).text = '1. Teacher(s):'
    table.cell(0, 1).merge(table.cell(0, 3)).text = general.get('teachers', '_______________')
    
    # Row 2: Grade + CEFR
    table.cell(1, 0).text = '2. Grade:'
    table.cell(1, 1).text = grade_labels.get(general.get('grade', ''), general.get('grade', ''))
    table.cell(1, 2).text = '3. CEFR Level:'
    table.cell(1, 3).text = general.get('cefr_level', '______')
    
    # Row 3: Trimester + Hours
    table.cell(2, 0).text = '4. Trimester:'
    table.cell(2, 1).text = general.get('trimester', '______')
    table.cell(2, 2).text = '5. Weekly Hour(s):'
    table.cell(2, 3).text = general.get('weekly_hours', '______')
    
    # Row 4: Week range
    table.cell(3, 0).text = '6. Week(s):'
    table.cell(3, 1).merge(table.cell(3, 3)).text = general.get('week_range', 'From week ____ to week ____')
    
    # Row 5: Scenario + Theme
    table.cell(4, 0).text = '7. Scenario __:'
    table.cell(4, 1).text = general.get('scenario', '')
    table.cell(4, 2).text = '8. Theme __:'
    table.cell(4, 3).text = general.get('theme', '')
    
    # Make first column bold
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Espacio


def add_standards_table(doc, standards):
    """Tabla de standards y learning outcomes"""
    skills = ['listening', 'reading', 'speaking', 'writing', 'mediation']
    labels = ['1. Listening:', '2. Reading:', '3. Speaking:', '4. Writing:', '5. Mediation:']
    
    table = doc.add_table(rows=len(skills) + 1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Skills:'
    hdr_cells[1].text = 'Specific Standards:'
    hdr_cells[2].text = 'Learning Outcomes:'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Rows
    for idx, skill in enumerate(skills):
        row = table.rows[idx + 1]
        row.cells[0].text = labels[idx]
        row.cells[0].paragraphs[0].runs[0].bold = True
        
        skill_data = standards.get(skill, {})
        
        # Standards
        std_list = []
        if isinstance(skill_data.get('specific'), list):
            std_list = skill_data['specific']
        elif isinstance(skill_data.get('specific_standards'), list):
            std_list = skill_data['specific_standards']
        
        if std_list:
            row.cells[1].text = '\n'.join([f"• {s}" for s in std_list])
        else:
            row.cells[1].text = 'To be defined'
        
        # Outcomes
        outcomes = skill_data.get('learning_outcomes', [])
        if outcomes:
            row.cells[2].text = '\n'.join([f"• {o}" for o in outcomes])
        else:
            row.cells[2].text = 'To be defined'
    
    doc.add_paragraph()


def add_competences_table(doc, competences):
    """Tabla de competencias comunicativas"""
    ling = competences.get('linguistic', {})
    prag = competences.get('pragmatic', '')
    socio = competences.get('sociolinguistic', '')
    
    grammar = ling.get('grammatical_features') or ling.get('grammar', [])
    vocab = ling.get('vocabulary', [])
    pronun = ling.get('phonemic_awareness') or ling.get('pronunciation', '')
    
    grammar_text = ', '.join(grammar) if isinstance(grammar, list) else str(grammar)
    vocab_text = ', '.join(vocab[:10]) if isinstance(vocab, list) else str(vocab)
    prag_text = ', '.join(prag) if isinstance(prag, list) else str(prag)
    socio_text = ', '.join(socio) if isinstance(socio, list) else str(socio)
    
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Linguistic Competence\n(Learn to Know)'
    hdr[1].text = 'Pragmatic Competence\n(Learn to Do)'
    hdr[2].text = 'Sociolinguistic Competence\n(Learn to Be)'
    for cell in hdr:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
    
    # Content
    cells = table.rows[1].cells
    
    ling_text = f"• Grammatical Features:\n{grammar_text or '______'}\n\n"
    ling_text += f"• Vocabulary:\n{vocab_text or '______'}\n\n"
    ling_text += f"• Pronunciation & Phonemic Awareness:\n{pronun or '______'}"
    cells[0].text = ling_text
    
    cells[1].text = prag_text or '______'
    cells[2].text = socio_text or '______'
    
    doc.add_paragraph()


def add_project_section(doc, project):
    """Sección de proyecto"""
    p = doc.add_paragraph()
    p.add_run('21st-Century Skills Project\n').bold = True
    p.add_run(f"Name: {project.get('name', '')}\n")
    p.add_run(f"Category: {project.get('category', '')}\n")
    p.add_run(f"Overview: {project.get('overview', '')}")
    doc.add_paragraph()


def add_objectives_table(doc, objectives):
    """Tabla de objetivos específicos"""
    skills = ['listening', 'reading', 'speaking', 'writing', 'mediation']
    labels = ['Listening', 'Reading', 'Speaking', 'Writing', 'Mediation']
    
    table = doc.add_table(rows=len(skills), cols=2)
    table.style = 'Table Grid'
    
    for idx, skill in enumerate(skills):
        row = table.rows[idx]
        row.cells[0].text = labels[idx]
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = objectives.get(skill, 'To be completed')
    
    doc.add_paragraph()


def add_materials_table(doc, materials):
    """Tabla de materiales"""
    mat_list = materials.get('required_materials', [])
    diff = materials.get('differentiated_instruction', 'To be completed')
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = 'Materials'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    
    if mat_list:
        table.cell(0, 1).text = '\n'.join([f"• {m}" for m in mat_list])
    else:
        table.cell(0, 1).text = 'To be completed'
    
    table.cell(1, 0).text = 'Differentiated Instruction and Accommodations for Students with Diverse Learning Needs (DLN)'
    table.cell(1, 0).paragraphs[0].runs[0].bold = True
    table.cell(1, 1).text = diff
    
    doc.add_paragraph()


def add_learning_sequence_table(doc, lessons, learning_seq):
    """Tabla de secuencia de aprendizaje"""
    dates = learning_seq.get('lesson_dates', ['', '', '', '', ''])
    
    table = doc.add_table(rows=len(lessons) + 1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Lesson'
    hdr[1].text = 'Learning Sequence'
    hdr[2].text = 'Date'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    
    # Lessons
    for idx, lesson in enumerate(lessons):
        row = table.rows[idx + 1]
        row.cells[0].text = f"Lesson {lesson.get('lesson_number', idx+1)}"
        row.cells[0].paragraphs[0].runs[0].bold = True
        
        seq_text = f"{lesson.get('skill_focus', '')}\n{lesson.get('specific_objective', '')}"
        row.cells[1].text = seq_text
        row.cells[2].text = dates[idx] if idx < len(dates) else '______'
    
    doc.add_paragraph()


def add_lesson_planners(doc, data):
    """Agregar 5 lesson planners"""
    lessons = data.get('lesson_planners', [])
    grade = data.get('grade', '')
    
    grade_labels = {
        'pre_k': 'Pre-K', 'K': 'Kindergarten',
        '1': 'Grade 1', '2': 'Grade 2', '3': 'Grade 3', '4': 'Grade 4',
        '5': 'Grade 5', '6': 'Grade 6', '7': 'Grade 7', '8': 'Grade 8',
        '9': 'Grade 9', '10': 'Grade 10', '11': 'Grade 11', '12': 'Grade 12',
    }
    
    for idx, lesson in enumerate(lessons):
        if idx > 0:
            doc.add_page_break()
        
        # Título
        title = doc.add_heading(f'Lesson Planner - Theme #___ - Lesson #{lesson.get("lesson_number", idx+1)}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tabla de info
        table = doc.add_table(rows=5, cols=6)
        table.style = 'Table Grid'
        
        # Row 1
        table.cell(0, 0).merge(table.cell(0, 2)).text = f"Lesson # {lesson.get('lesson_number', idx+1)}"
        table.cell(0, 3).merge(table.cell(0, 5)).text = f"Skills focus: {lesson.get('skill_focus', '')}"
        
        # Row 2
        table.cell(1, 0).text = 'Grade:'
        table.cell(1, 1).text = grade_labels.get(grade, grade)
        table.cell(1, 2).text = 'Scenario:'
        table.cell(1, 3).text = lesson.get('scenario', '')
        table.cell(1, 4).text = 'Theme:'
        table.cell(1, 5).text = lesson.get('theme', '')
        
        # Row 3
        table.cell(2, 0).text = 'Date(s):'
        table.cell(2, 1).merge(table.cell(2, 2)).text = lesson.get('date', '______')
        table.cell(2, 3).text = 'Time:'
        table.cell(2, 4).merge(table.cell(2, 5)).text = lesson.get('time', '45-60 minutes')
        
        # Row 4
        table.cell(3, 0).text = 'Specific Objective:'
        table.cell(3, 1).merge(table.cell(3, 5)).text = lesson.get('specific_objective', '______')
        
        # Row 5
        table.cell(4, 0).text = 'Learning Outcome:'
        table.cell(4, 1).merge(table.cell(4, 5)).text = lesson.get('learning_outcome', '______')
        
        doc.add_paragraph()
        
        # Stages
        doc.add_heading('The Six Action-oriented Approach Lesson Stages', level=2)
        add_stages_table(doc, lesson.get('lesson_stages', []))
        
        # Comments
        doc.add_heading('Comments and Observations', level=2)
        add_comments_table(doc, lesson.get('comments', {}))


def add_stages_table(doc, stages):
    """Tabla de stages"""
    stage_names = [
        'Stage 1 - Warm-up / Pre-task (Engagement, Modeling and Clarification):',
        'Stage 2 - Presentation:',
        'Stage 3 - Preparation:',
        'Stage 4 - Performance:',
        'Stage 5 - Assessment / Post-task:',
        'Stage 6 - Reflection:',
    ]
    
    table = doc.add_table(rows=len(stages) + 1, cols=2)
    table.style = 'Table Grid'
    
    # Header
    table.rows[0].cells[0].text = 'Stage'
    table.rows[0].cells[1].text = 'Time'
    
    for idx, stage in enumerate(stages):
        row = table.rows[idx + 1]
        
        # Stage content
        activities = stage.get('activities', [])
        stage_text = stage_names[idx] + '\n'
        stage_text += '\n'.join([f"• {act}" for act in activities])
        row.cells[0].text = stage_text
        
        # Time
        row.cells[1].text = stage.get('estimated_time', '___')
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()


def add_comments_table(doc, comments):
    """Tabla de comentarios"""
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = 'Homework:'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    table.cell(0, 1).text = comments.get('homework', '______')
    
    table.cell(1, 0).text = 'Formative Assessment:'
    table.cell(1, 0).paragraphs[0].runs[0].bold = True
    table.cell(1, 1).text = comments.get('formative_assessment', '______')
    
    table.cell(2, 0).text = "Teacher's Comments:"
    table.cell(2, 0).paragraphs[0].runs[0].bold = True
    table.cell(2, 1).text = comments.get('teacher_comments', '______')
    
    doc.add_paragraph()
